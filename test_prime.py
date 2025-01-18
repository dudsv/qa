import os
import requests
from tkinter import Tk, Text, Button, Label, filedialog, messagebox, Frame
from tkinter.ttk import Notebook
from urllib.parse import urlparse
from bs4 import BeautifulSoup
import docx
import pandas as pd
from difflib import SequenceMatcher
from openpyxl import load_workbook
from openpyxl.styles import PatternFill
from docx.shared import RGBColor
import openpyxl
import difflib

# Função para buscar URLs a partir do sitemap
def fetch_urls_from_sitemap(sitemap_url):
    """Fetch URLs from a sitemap."""
    try:
        response = requests.get(sitemap_url)
        response.raise_for_status()
        soup = BeautifulSoup(response.content, 'xml')
        urls = [loc.text for loc in soup.find_all('loc')]
        return urls
    except Exception as e:
        messagebox.showerror("Erro", f"Erro ao processar o sitemap: {e}")
        return []

# Função para salvar conteúdo da URL em um documento .docx
# Função para avaliar se o conteúdo de um .docx está presente em uma URL
def evaluate_doc_against_url():
    """Evaluate if the content of a .docx is fully present in a URL."""
    doc_path = filedialog.askopenfilename(title="Selecione o arquivo .docx", filetypes=[("Documentos Word", "*.docx")])
    if not doc_path:
        return

    url = url_entry.get("1.0", "end").strip()
    if not url:
        messagebox.showwarning("Aviso", "Por favor, insira uma URL para avaliação.")
        return

    try:
        # Extract text from the .docx file
        doc = docx.Document(doc_path)
        doc_text_lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        # Fetch text from the URL
        response = requests.get(url)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, 'html.parser')
        page_text_lines = [tag.get_text(strip=True) for tag in soup.find_all(['h1', 'h2', 'h3', 'p'])]

        # Prepare the comparison results
        results = []
        for doc_line in doc_text_lines:
            best_match = max(page_text_lines, key=lambda x: SequenceMatcher(None, doc_line, x).ratio())
            similarity = SequenceMatcher(None, doc_line, best_match).ratio()

            if similarity == 1.0:
                status = "Presente"
            elif similarity >= 0.5:
                status = "Semelhante"
            else:
                status = "Ausente"

            # Find differences using difflib
            diff = list(SequenceMatcher(None, doc_line, best_match).get_opcodes())
            diff_output = ""
            for tag, i1, i2, j1, j2 in diff:
                if tag == 'replace' or tag == 'delete':
                    diff_output += f"Discrepância: {doc_line[i1:i2]} → {best_match[j1:j2]}\n"

            results.append({
                "Texto do Doc": doc_line,
                "Texto da Página": best_match if similarity > 0 else "N/A",
                "Status": status,
                "Discrepâncias": diff_output if diff_output else "Nenhuma"
            })

        # Seleção da pasta de saída
        output_file = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel Files", "*.xlsx")], title="Salvar Resultado de Avaliação")
        if not output_file:
            return

        # Save the results to an Excel file
        df = pd.DataFrame(results)
        df.to_excel(output_file, index=False)

        # Load the Excel file for styling
        wb = load_workbook(output_file)
        ws = wb.active

        # Define color fills
        green_fill = PatternFill(start_color="00FF00", end_color="00FF00", fill_type="solid")
        yellow_fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
        red_fill = PatternFill(start_color="FF0000", end_color="FF0000", fill_type="solid")

        # Apply colors to the "Status" column based on the value
        for row in range(2, len(results) + 2):  # Starting from row 2 to avoid the header
            status_cell = ws.cell(row=row, column=3)
            if status_cell.value == "Presente":
                status_cell.fill = green_fill
            elif status_cell.value == "Semelhante":
                status_cell.fill = yellow_fill
            elif status_cell.value == "Ausente":
                status_cell.fill = red_fill

        # Save the styled Excel file
        wb.save(output_file)

        messagebox.showinfo("Concluído", f"Avaliação concluída! Arquivo salvo em:\n{output_file}")
    except Exception as e:
        messagebox.showerror("Erro", f"Erro ao avaliar o conteúdo: {e}")

# Função para salvar conteúdo das URLs em uma planilha Excel
def save_content_to_excel(urls, folder):
    """Fetch content from URLs and save to an Excel file."""
    data = []
    for url in urls:
        try:
            response = requests.get(url)
            response.raise_for_status()
            soup = BeautifulSoup(response.text, 'html.parser')
            text_content = "\n".join([tag.get_text(strip=True) for tag in soup.find_all(['h1', 'h2', 'h3', 'p'])])
            data.append({"URL": url, "Conteúdo": text_content})
        except Exception as e:
            data.append({"URL": url, "Conteúdo": f"Erro: {e}"})

    file_path = os.path.join(folder, "conteudo_urls.xlsx")
    df = pd.DataFrame(data)
    df.to_excel(file_path, index=False)
    return file_path

# Função para processar URLs e salvar conteúdo
def process_urls(urls, output_dir, save_as_doc):
    """Process and save content for each URL."""
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    if save_as_doc:
        for url in urls:
            url = url.strip()
            if url:
                status = save_content_to_doc(url, output_dir)
                print(f"Processed {url}: {status}")
    else:
        file_path = save_content_to_excel(urls, output_dir)
        print(f"Saved content to Excel: {file_path}")

# Função para iniciar o scraper
def start_scraper():
    """Start the scraper with input URLs or sitemap."""
    input_text = url_input.get("1.0", "end").strip()
    urls = []

    if input_text.startswith("http") and input_text.endswith(".xml"):
        urls = fetch_urls_from_sitemap(input_text)
    else:
        urls = [line.strip() for line in input_text.splitlines() if line.strip()]

    if not urls:
        messagebox.showwarning("Aviso", "Nenhuma URL fornecida ou válida encontrada.")
        return

    output_dir = filedialog.askdirectory(title="Selecione o diretório de saída")
    if not output_dir:
        return

    save_as_doc = messagebox.askyesno("Salvar como", "Deseja salvar como documentos Word (.docx)? (Selecionar 'Não' salvará como Excel)")
    process_urls(urls, output_dir, save_as_doc)
    messagebox.showinfo("Concluído", f"Conteúdo salvo na pasta: {output_dir}")

# Função para avaliar se o conteúdo de um .docx está presente em uma URL
def evaluate_doc_against_url():
    """Evaluate if the content of a .docx is fully present in a URL."""
    doc_path = filedialog.askopenfilename(title="Selecione o arquivo .docx", filetypes=[("Documentos Word", "*.docx")])
    if not doc_path:
        return

    url = url_entry.get("1.0", "end").strip()
    if not url:
        messagebox.showwarning("Aviso", "Por favor, insira uma URL para avaliação.")
        return

    try:
        # Extract text from the .docx file
        doc = docx.Document(doc_path)
        doc_text_lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        # Fetch text from the URL
        response = requests.get(url)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, 'html.parser')
        page_text_lines = [tag.get_text(strip=True) for tag in soup.find_all(['h1', 'h2', 'h3', 'p'])]

        # Prepare the comparison results
        results = []
        for doc_line in doc_text_lines:
            best_match = max(page_text_lines, key=lambda x: SequenceMatcher(None, doc_line, x).ratio())
            similarity = SequenceMatcher(None, doc_line, best_match).ratio()

            if similarity == 1.0:
                status = "Presente"
            elif similarity >= 0.5:
                status = "Semelhante"
            else:
                status = "Ausente"

            results.append({
                "Texto do Doc": doc_line,
                "Texto da Página": best_match if similarity > 0 else "N/A",
                "Status": status
            })

        # Seleção da pasta de saída
        output_file = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel Files", "*.xlsx")], title="Salvar Resultado de Avaliação")
        if not output_file:
            return

        # Save the results to an Excel file
        df = pd.DataFrame(results)
        df.to_excel(output_file, index=False)

        # Load the Excel file for styling
        wb = load_workbook(output_file)
        ws = wb.active

        # Define color fills
        green_fill = PatternFill(start_color="00FF00", end_color="00FF00", fill_type="solid")
        yellow_fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
        red_fill = PatternFill(start_color="FF0000", end_color="FF0000", fill_type="solid")

        # Apply colors to the "Status" column based on the value
        for row in range(2, len(results) + 2):  # Starting from row 2 to avoid the header
            status_cell = ws.cell(row=row, column=3)
            if status_cell.value == "Presente":
                status_cell.fill = green_fill
            elif status_cell.value == "Semelhante":
                status_cell.fill = yellow_fill
            elif status_cell.value == "Ausente":
                status_cell.fill = red_fill

        # Save the styled Excel file
        wb.save(output_file)

        messagebox.showinfo("Concluído", f"Avaliação concluída! Arquivo salvo em:\n{output_file}")
    except Exception as e:
        messagebox.showerror("Erro", f"Erro ao avaliar o conteúdo: {e}")

# Configuração da UI
root = Tk()
root.title("Web Scraper e Avaliador")
root.geometry("600x500")

notebook = Notebook(root)
notebook.pack(fill="both", expand=True)

# Aba 1: Web Scraper
scraper_tab = Frame(notebook)
notebook.add(scraper_tab, text="Web Scraper")

Label(scraper_tab, text="Insira URLs (uma por linha) ou o URL de um sitemap:").pack(pady=10)
url_input = Text(scraper_tab, height=15, width=70)
url_input.pack(pady=10)
Button(scraper_tab, text="Iniciar Scraper", command=start_scraper, bg="purple", fg="white").pack(pady=10)

# Aba 2: Avaliador de Conteúdo
evaluator_frame = Frame(notebook)
notebook.add(evaluator_frame, text="Avaliador de Conteúdo")

Label(evaluator_frame, text="Insira a URL para avaliação:").pack(pady=10)
url_entry = Text(evaluator_frame, height=1, width=70)
url_entry.pack(pady=10)
Button(evaluator_frame, text="Iniciar Avaliação", command=evaluate_doc_against_url, bg="purple", fg="white").pack(pady=10)

root.mainloop()
